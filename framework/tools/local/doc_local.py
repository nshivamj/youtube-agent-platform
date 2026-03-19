"""Document generation tools — creates Word docs for reports and summaries."""

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt

_REPORTS_DIR = Path(__file__).parent.parent.parent.parent / "reports"


class DocLocalTools:
    name = "doc_tools"
    description = "Generate Word document reports and summaries"
    tool_names = [
        "generate_summary_doc",
    ]

    def __init__(self):
        _REPORTS_DIR.mkdir(exist_ok=True)

    async def generate_summary_doc(self, title: str, summary: str, extra_content: str = "") -> dict:
        """Generate a Word document with a project summary.

        Args:
            title: Document title (e.g. 'Project Summary - PROJ').
            summary: Main summary text to include in the document.
            extra_content: Optional additional content in markdown-like format.
                           Use '## Heading' for section headings and plain text for content.
                           Example: '## Commits\\nCommit 1...\\n## Copilot Usage\\n3 out of 5 used copilot'

        Returns:
            Dict with file_path and success status.
        """
        doc = Document()

        # Title
        doc.add_heading(title, level=0)

        # Metadata
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph("")

        # Summary section
        doc.add_heading("Summary", level=1)
        doc.add_paragraph(summary)

        # Parse extra content — lines starting with '## ' become headings
        if extra_content:
            for line in extra_content.split("\n"):
                stripped = line.strip()
                if stripped.startswith("## "):
                    doc.add_heading(stripped[3:], level=1)
                elif stripped:
                    doc.add_paragraph(stripped)

        # Style the document
        style = doc.styles["Normal"]
        style.font.size = Pt(11)
        style.font.name = "Calibri"

        # Save
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_title = title.lower().replace(" ", "_").replace("-", "_")[:30]
        filename = f"{safe_title}_{timestamp}.docx"
        file_path = _REPORTS_DIR / filename
        doc.save(str(file_path))

        return {
            "file_path": str(file_path),
            "filename": filename,
            "success": True,
        }
